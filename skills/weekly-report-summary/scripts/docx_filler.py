# -*- coding: utf-8 -*-
"""
Word 文档填充模块

读取模板并填充周报内容

模板结构说明：
- 整个文档是一个大表格
- R0: 星光闪烁标题
- R1: 点赞内容
- R2: 诚信坦荡标题  
- R3-R9: 建议内容
- R10: 工作复盘与规划标题
- R11: 所有人的周报内容（在一个大单元格里）
"""

import os
import re
import glob
from copy import deepcopy
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

import config


def create_shading_element(fill_color: str) -> OxmlElement:
    """
    创建底纹元素（用于设置文字背景色/高亮）
    
    Args:
        fill_color: 十六进制颜色值，如 "FFFF00" 表示黄色
    """
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    return shd


def copy_run_format(source_run, target_run):
    """
    复制 run 的格式属性（字体、字号、颜色、底纹等）
    """
    # 复制基本格式
    if source_run.font.name:
        target_run.font.name = source_run.font.name
    if source_run.font.size:
        target_run.font.size = source_run.font.size
    if source_run.font.bold is not None:
        target_run.font.bold = source_run.font.bold
    if source_run.font.italic is not None:
        target_run.font.italic = source_run.font.italic
    
    # 复制东亚字体
    source_rPr = source_run._element.rPr
    if source_rPr is not None:
        rFonts = source_rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            eastAsia = rFonts.get(qn('w:eastAsia'))
            if eastAsia:
                target_run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), eastAsia)
        
        # 复制字色
        color = source_rPr.find(qn('w:color'))
        if color is not None:
            target_color = OxmlElement('w:color')
            target_color.set(qn('w:val'), color.get(qn('w:val')))
            target_rPr = target_run._element.get_or_add_rPr()
            # 移除已有的 color
            existing_color = target_rPr.find(qn('w:color'))
            if existing_color is not None:
                target_rPr.remove(existing_color)
            target_rPr.append(target_color)
        
        # 复制字号
        sz = source_rPr.find(qn('w:sz'))
        if sz is not None:
            target_sz = OxmlElement('w:sz')
            target_sz.set(qn('w:val'), sz.get(qn('w:val')))
            target_rPr = target_run._element.get_or_add_rPr()
            existing_sz = target_rPr.find(qn('w:sz'))
            if existing_sz is not None:
                target_rPr.remove(existing_sz)
            target_rPr.append(target_sz)


def apply_shading_to_run(run, fill_color: str):
    """
    为 run 应用底纹颜色
    """
    rPr = run._element.get_or_add_rPr()
    # 移除已有的 shd
    existing_shd = rPr.find(qn('w:shd'))
    if existing_shd is not None:
        rPr.remove(existing_shd)
    # 添加新的 shd
    rPr.append(create_shading_element(fill_color))


def copy_paragraph_format(source_para, target_para):
    """
    复制段落格式（行距、对齐等）
    """
    source_pPr = source_para._element.pPr
    if source_pPr is not None:
        target_pPr = target_para._element.get_or_add_pPr()
        
        # 复制行距
        spacing = source_pPr.find(qn('w:spacing'))
        if spacing is not None:
            target_spacing = OxmlElement('w:spacing')
            for attr in spacing.attrib:
                target_spacing.set(attr, spacing.get(attr))
            existing_spacing = target_pPr.find(qn('w:spacing'))
            if existing_spacing is not None:
                target_pPr.remove(existing_spacing)
            target_pPr.append(target_spacing)
        
        # 复制对齐
        jc = source_pPr.find(qn('w:jc'))
        if jc is not None:
            target_jc = OxmlElement('w:jc')
            target_jc.set(qn('w:val'), jc.get(qn('w:val')))
            existing_jc = target_pPr.find(qn('w:jc'))
            if existing_jc is not None:
                target_pPr.remove(existing_jc)
            target_pPr.append(target_jc)


def apply_saved_run_format(run, saved_rPr_xml):
    """
    从保存的 rPr XML 副本应用 run 格式
    """
    if saved_rPr_xml is None:
        return
    
    rPr = run._element.get_or_add_rPr()
    
    # 复制字色
    color = saved_rPr_xml.find(qn('w:color'))
    if color is not None:
        existing = rPr.find(qn('w:color'))
        if existing is not None:
            rPr.remove(existing)
        rPr.append(deepcopy(color))
    
    # 复制字号
    sz = saved_rPr_xml.find(qn('w:sz'))
    if sz is not None:
        existing = rPr.find(qn('w:sz'))
        if existing is not None:
            rPr.remove(existing)
        rPr.append(deepcopy(sz))
    
    # 复制字体
    rFonts = saved_rPr_xml.find(qn('w:rFonts'))
    if rFonts is not None:
        existing = rPr.find(qn('w:rFonts'))
        if existing is not None:
            rPr.remove(existing)
        rPr.append(deepcopy(rFonts))


def apply_saved_para_format(para, saved_pPr_xml):
    """
    从保存的 pPr XML 副本应用段落格式
    """
    if saved_pPr_xml is None:
        return
    
    pPr = para._element.get_or_add_pPr()
    
    # 复制行距
    spacing = saved_pPr_xml.find(qn('w:spacing'))
    if spacing is not None:
        existing = pPr.find(qn('w:spacing'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(deepcopy(spacing))
    
    # 复制对齐
    jc = saved_pPr_xml.find(qn('w:jc'))
    if jc is not None:
        existing = pPr.find(qn('w:jc'))
        if existing is not None:
            pPr.remove(existing)
        pPr.append(deepcopy(jc))


def find_template_file() -> str:
    """查找模板文件"""
    # 首先尝试配置的路径
    if os.path.exists(config.TEMPLATE_PATH):
        return config.TEMPLATE_PATH
    
    # 再尝试 skill 内常见模板目录
    skill_dir = config.SKILL_DIR
    template_files = (
        glob.glob(os.path.join(skill_dir, "templates", "*.docx"))
        + glob.glob(os.path.join(skill_dir, "assets", "*.docx"))
        + glob.glob(os.path.join(skill_dir, "*.docx"))
    )
    
    if template_files:
        return template_files[0]
    
    raise FileNotFoundError("未找到模板文件")


def analyze_work_content_cell(cell) -> Dict[str, int]:
    """
    分析工作内容单元格中各成员的位置
    
    Returns:
        {花名: 段落索引}
    """
    member_positions = {}
    
    for i, para in enumerate(cell.paragraphs):
        text = para.text.strip()
        
        # 检查是否是花名（在成员列表中）
        for member in config.MEMBER_TO_GROUP.keys():
            if text == member:
                member_positions[member] = i
                break
        
        # 也检查组名
        for group in config.GROUP_MEMBERS.keys():
            if text == group:
                member_positions[f"__group__{group}"] = i
    
    return member_positions


def get_member_content_range(cell, member_positions: Dict, member_name: str) -> Tuple[int, int]:
    """
    获取某个成员在单元格中的内容范围
    
    Returns:
        (start_idx, end_idx) 段落索引范围
    """
    if member_name not in member_positions:
        return (-1, -1)
    
    start_idx = member_positions[member_name]
    
    # 找到下一个成员或组的位置作为结束
    end_idx = len(cell.paragraphs)
    
    for name, idx in member_positions.items():
        if idx > start_idx and idx < end_idx:
            end_idx = idx
    
    return (start_idx, end_idx)


def update_member_in_cell(cell, start_idx: int, content: Dict, member_positions: Dict) -> bool:
    """
    更新单元格中某个成员的周报内容
    
    Args:
        cell: 表格单元格
        start_idx: 花名所在段落索引
        content: 解析后的周报内容
        member_positions: 所有成员位置字典（只包含配置中的成员）
        
    Returns:
        是否成功更新
    """
    paragraphs = cell.paragraphs
    
    # 找到该成员内容的结束位置
    # 注意：需要检测下一个"本周完成工作"来判断另一个人的周报开始
    # 因为模板中可能有不在配置中的成员
    end_idx = len(paragraphs)
    
    # 首先用配置中的成员位置限制范围
    for name, idx in member_positions.items():
        if idx > start_idx and idx < end_idx:
            end_idx = idx
    
    # 定义部分映射（按长度从长到短排列，避免短标签误匹配长标签）
    # 例如："本周得与失" 应该匹配 gains_losses，而不是被 "本周" 匹配到 this_week_work
    sections_map_ordered = [
        # 长标签优先（必须先匹配，避免被短标签截胡）
        ('本周主要工作项', 'this_week_work'),
        ('本周完成工作', 'this_week_work'),
        ('本周工作回顾', 'this_week_work'),
        ('下周主要计划工作项', 'next_week_plan'),
        ('下周工作计划', 'next_week_plan'),
        ('下周完成工作', 'next_week_plan'),
        ('其他收获与思考', 'gains_losses'),
        ('本周得与失', 'gains_losses'),
        ('得与失', 'gains_losses'),
        # 中等长度标签
        ('本周完成', 'this_week_work'),
        ('本周工作', 'this_week_work'),
        ('下周工作', 'next_week_plan'),
        ('下周计划', 'next_week_plan'),
        # 短标签放最后（只有精确匹配或后面跟冒号/空格才算）
        ('本周', 'this_week_work'),
        ('下周', 'next_week_plan'),
    ]
    
    # 转换为字典供后续使用
    sections_map = dict(sections_map_ordered)
    
    # 第一遍扫描：找到各部分的位置
    # 只取第一次出现的每个部分（避免跨越到下一个人的周报）
    section_ranges = {}  # {section_key: (header_idx, content_start, content_end)}
    current_section = None
    section_header_idx = None
    seen_sections = set()  # 已经看到过的部分
    this_week_work_count = 0  # 计数"本周完成工作"出现次数
    
    for i in range(start_idx + 1, end_idx):
        para = paragraphs[i]
        text = para.text.strip()
        
        # 检查是否是部分标题（使用有序列表，长标签优先匹配）
        matched_section = None
        for label, key in sections_map_ordered:
            # 短标签（2个字）需要更精确的匹配，避免误匹配
            if len(label) <= 2:
                # 精确匹配，或者后面紧跟冒号/数字
                if text == label or text.startswith(label + '：') or text.startswith(label + ':'):
                    matched_section = key
                    break
            else:
                # 长标签可以用 startswith
                if text == label or text.startswith(label):
                    matched_section = key
                    break
        
        if matched_section:
            # 如果再次遇到"本周完成工作"，说明进入了下一个人的周报，停止扫描
            if matched_section == 'this_week_work':
                this_week_work_count += 1
                if this_week_work_count > 1:
                    # 保存当前部分并退出
                    if current_section and section_header_idx is not None:
                        section_ranges[current_section] = (section_header_idx, section_header_idx + 1, i)
                    break
            
            # 保存上一个部分的范围
            if current_section and section_header_idx is not None:
                section_ranges[current_section] = (section_header_idx, section_header_idx + 1, i)
            
            # 只记录第一次出现的部分
            if matched_section not in seen_sections:
                current_section = matched_section
                section_header_idx = i
                seen_sections.add(matched_section)
            else:
                # 重复的部分，跳过
                current_section = None
                section_header_idx = None
    else:
        # 正常结束循环，保存最后一个部分
        if current_section and section_header_idx is not None:
            section_ranges[current_section] = (section_header_idx, section_header_idx + 1, end_idx)
    
    # 第二遍：更新各部分内容
    updated_sections = set()
    
    for section_key, (header_idx, content_start, content_end) in section_ranges.items():
        new_content = content.get(section_key, '')
        if not new_content:
            continue
        
        # 清除旧内容段落（从 content_start 到 content_end-1）
        first_content_para = None
        for i in range(content_start, content_end):
            para = paragraphs[i]
            text = para.text.strip()
            
            # 跳过空段落
            if not text:
                continue
            
            if first_content_para is None:
                # 第一个内容段落：更新为新内容
                first_content_para = para
                update_paragraph_content(para, new_content)
                updated_sections.add(section_key)
            else:
                # 后续内容段落：清空（这些是模板中的占位内容）
                para.clear()
    
    return len(updated_sections) > 0


def remove_paragraph_numbering(para):
    """
    移除段落的Word自动编号属性
    """
    pPr = para._p.pPr
    if pPr is not None:
        numPr = pPr.find(qn('w:numPr'))
        if numPr is not None:
            pPr.remove(numPr)


def remove_paragraph(para):
    """
    从文档中删除段落
    """
    p = para._element
    p.getparent().remove(p)


def clean_empty_paragraphs(cell, member_names: set = None):
    """
    清理单元格中的空段落，但在成员之间保留一个空行分隔
    
    Args:
        cell: 单元格
        member_names: 成员花名集合，用于识别成员边界
    """
    if member_names is None:
        member_names = set()
    
    paragraphs = cell.paragraphs
    removed_count = 0
    
    # 收集需要删除的空段落
    to_remove = []
    prev_was_empty = False
    
    for i, para in enumerate(paragraphs):
        text = para.text.strip()
        is_empty = not text
        is_member_name = text in member_names
        
        if is_empty:
            # 检查下一个段落是否是成员名（需要保留空行作为分隔）
            next_is_member = False
            if i + 1 < len(paragraphs):
                next_text = paragraphs[i + 1].text.strip()
                next_is_member = next_text in member_names
            
            if next_is_member and not prev_was_empty:
                # 保留这个空行作为成员之间的分隔
                prev_was_empty = True
            else:
                # 删除这个空行
                to_remove.append(para)
        else:
            prev_was_empty = False
    
    # 从后往前删除
    for para in reversed(to_remove):
        remove_paragraph(para)
        removed_count += 1
    
    return removed_count


def update_paragraph_content(para, new_text: str):
    """
    更新段落内容，保持格式，并移除编号属性
    """
    if not new_text:
        return
    
    # 去除内容末尾的空白和换行，避免段落末尾多余空行
    new_text = new_text.rstrip()
    
    # 获取原格式
    original_runs = para.runs
    font_name = None
    font_size = None
    
    if original_runs:
        first_run = original_runs[0]
        font_name = first_run.font.name
        font_size = first_run.font.size
    
    # 清除原内容
    para.clear()
    
    # 移除编号属性（避免Word自动编号残留）
    remove_paragraph_numbering(para)
    
    # 添加新内容
    run = para.add_run(new_text)
    
    # 应用格式
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if font_size:
        run.font.size = font_size


def fill_praise_section(table, organized_reports: Dict[str, Dict]):
    """
    填充推车（点赞）区域
    
    模板结构：R1 包含 "关键词" 和 "事迹推车" 标题，需要在后面添加点赞内容
    格式要求：
    - 人名段落：黄色底纹（FFFF00），与模版示例格式一致
    - 内容段落：无底纹，与模版示例格式一致
    """
    from content_parser import is_praise_template_content
    
    # 收集所有人的点赞内容
    all_praises = []
    for name, content in organized_reports.items():
        praise = content.get('praise', '')
        if praise and not is_template_content(praise):
            # 进一步过滤，检查每一行是否是模板内容
            praise_lines = praise.split('\n')
            valid_lines = [line for line in praise_lines 
                          if line.strip() and not is_praise_template_content(line.strip())]
            if valid_lines:
                all_praises.append({
                    'name': name,
                    'content': '\n'.join(valid_lines)
                })
    
    if not all_praises:
        print("未找到点赞内容")
        return
    
    print(f"收集到 {len(all_praises)} 人的点赞内容")
    
    # 找到 R1（推车区域）
    if len(table.rows) < 2:
        return
    
    praise_cell = table.rows[1].cells[1]  # 第二列是内容区域
    
    # 保存模版中的格式样本（在清除前提取）
    template_name_para = None  # 人名段落样本
    template_content_para = None  # 内容段落样本
    template_name_run = None
    template_content_run = None
    
    # 深拷贝格式信息，因为后面会删除原段落
    template_name_rPr_xml = None
    template_content_rPr_xml = None
    template_pPr_xml = None
    
    if len(praise_cell.paragraphs) >= 2:
        # 第一个段落是人名示例（如"人员1"），有黄色底纹
        template_name_para = praise_cell.paragraphs[0]
        if template_name_para.runs:
            template_name_run = template_name_para.runs[0]
            # 保存格式 XML 副本
            if template_name_run._element.rPr is not None:
                template_name_rPr_xml = deepcopy(template_name_run._element.rPr)
        if template_name_para._element.pPr is not None:
            template_pPr_xml = deepcopy(template_name_para._element.pPr)
        
        # 第二个段落是内容示例（如"XXXXXX"），无底纹
        template_content_para = praise_cell.paragraphs[1]
        if template_content_para.runs:
            template_content_run = template_content_para.runs[0]
            if template_content_run._element.rPr is not None:
                template_content_rPr_xml = deepcopy(template_content_run._element.rPr)
    
    # 删除所有现有段落（保留第一个，因为单元格必须有至少一个段落）
    # 从后往前删除以避免索引问题
    paragraphs_to_remove = list(praise_cell.paragraphs[1:])
    for para in reversed(paragraphs_to_remove):
        remove_paragraph(para)
    
    # 清除第一个段落的内容
    if praise_cell.paragraphs:
        praise_cell.paragraphs[0].clear()
    
    # 为每个人添加单独的段落，并应用模版格式
    first_para = True
    for i, item in enumerate(all_praises):
        name = item['name']
        content = item['content']
        
        # === 人名段落 ===
        if first_para and praise_cell.paragraphs:
            name_para = praise_cell.paragraphs[0]
            first_para = False
        else:
            name_para = praise_cell.add_paragraph()
        
        # 添加人名文本
        name_run = name_para.add_run(name)
        
        # 应用模版格式（使用保存的 XML 副本）
        apply_saved_run_format(name_run, template_name_rPr_xml)
        apply_saved_para_format(name_para, template_pPr_xml)
        
        # 应用黄色底纹（人名标黄）
        apply_shading_to_run(name_run, 'FFFF00')
        
        # === 内容段落 ===
        content_para = praise_cell.add_paragraph()
        content_run = content_para.add_run(content)
        
        # 应用模版格式（无底纹）
        apply_saved_run_format(content_run, template_content_rPr_xml)
        apply_saved_para_format(content_para, template_pPr_xml)
        
        # 空行分隔（最后一个人不需要）
        if i < len(all_praises) - 1:
            sep_para = praise_cell.add_paragraph()
            # 空行也应用模版段落格式
            apply_saved_para_format(sep_para, template_pPr_xml)


def is_praise_placeholder(text: str) -> bool:
    """检查是否是点赞区域的占位内容"""
    placeholders = ['人员1', '人员2', '人员3', 'XXXXXX', 'XXX']
    return text in placeholders


def fill_suggestion_section(table, organized_reports: Dict[str, Dict]):
    """
    填充建议（诚信坦荡）区域
    
    模板结构：
    - R3: 表头（方向, 建议人, @ta推进, 存在问题, 建设性意见与措施）
    - R4-R5: 降本增效
    - R6-R7: 协同攻坚
    - R8-R9: 高质量发展
    """
    from content_parser import extract_all_suggestions
    
    # 提取所有建议
    all_suggestions = extract_all_suggestions(organized_reports)
    
    if not all_suggestions:
        print("未找到建议内容")
        return
    
    print(f"收集到 {len(all_suggestions)} 条建议内容")
    
    # 按方向分类
    suggestions_by_direction = {
        '降本增效': [],
        '协同攻坚': [],
        '高质量发展': []
    }
    
    for suggestion in all_suggestions:
        direction = suggestion.get('direction', '')
        if direction in suggestions_by_direction:
            suggestions_by_direction[direction].append(suggestion)
        else:
            # 如果没有明确方向，放到降本增效
            suggestions_by_direction['降本增效'].append(suggestion)
    
    # 方向对应的行范围
    direction_rows = {
        '降本增效': [4, 5],
        '协同攻坚': [6, 7],
        '高质量发展': [8, 9]
    }
    
    # 填充每个方向的建议
    for direction, rows in direction_rows.items():
        suggestions = suggestions_by_direction.get(direction, [])
        
        for i, row_idx in enumerate(rows):
            if i >= len(suggestions):
                break
            
            if row_idx >= len(table.rows):
                break
            
            row = table.rows[row_idx]
            suggestion = suggestions[i]
            
            # 填充各列
            # C1: 建议人
            if len(row.cells) > 1:
                cell = row.cells[1]
                if cell.paragraphs:
                    update_paragraph_content(cell.paragraphs[0], suggestion.get('name', ''))
            
            # C2: @ta推进
            if len(row.cells) > 2:
                cell = row.cells[2]
                if cell.paragraphs:
                    update_paragraph_content(cell.paragraphs[0], suggestion.get('target', ''))
            
            # C3: 存在问题
            if len(row.cells) > 3:
                cell = row.cells[3]
                if cell.paragraphs:
                    update_paragraph_content(cell.paragraphs[0], suggestion.get('problem', ''))
            
            # C4: 建设性意见与措施
            if len(row.cells) > 4:
                cell = row.cells[4]
                if cell.paragraphs:
                    update_paragraph_content(cell.paragraphs[0], suggestion.get('suggestion', ''))


def is_template_content(text: str) -> bool:
    """
    检查是否是模板占位内容（不是真实的点赞/建议）
    """
    template_markers = [
        '事迹推车', '关键词', '建议人', '@ta推进', 
        '存在问题', '建设性意见与措施', '工作复盘',
        '降本增效', '协同攻坚', '高质量发展',
        '助我前行', '协同', '专业', '突破'
    ]
    
    # 如果内容主要是模板标记，认为是无效内容
    for marker in template_markers:
        if marker in text:
            return True
    
    return False


def remove_xxx_placeholders_and_numbering(cell):
    """
    移除单元格中所有的 XXX 占位内容，并移除所有空段落和内容段落的Word自动编号属性
    """
    removed_count = 0
    for para in cell.paragraphs:
        text = para.text.strip()
        
        # 清除包含XXX的段落
        if 'XXX' in text:
            para.clear()
            removed_count += 1
        
        # 移除所有段落的编号属性（避免残留单独的编号）
        # 包括空段落和有内容的段落
        remove_paragraph_numbering(para)
    
    return removed_count


def fill_template_by_replacement(
    organized_reports: Dict[str, Dict],
    template_path: str = None,
    output_path: str = None
) -> str:
    """
    通过直接替换模板中的内容来填充周报
    """
    template_path = template_path or find_template_file()
    output_path = output_path or config.OUTPUT_PATH
    
    print(f"使用模板: {template_path}")
    
    doc = Document(template_path)
    
    if not doc.tables:
        print("模板中没有表格，使用简单格式")
        return create_simple_report(organized_reports, output_path)
    
    table = doc.tables[0]
    
    # === 1. 填充推车（点赞）区域 R1 ===
    fill_praise_section(table, organized_reports)
    
    # === 1.5 填充建议（诚信坦荡）区域 R3-R9 ===
    fill_suggestion_section(table, organized_reports)
    
    # === 2. 找到工作复盘内容所在的单元格 ===
    # 模板结构：R10 是"工作复盘与规划"标题，R11 是实际内容
    work_cell = None
    work_row_idx = None
    for row_idx, row in enumerate(table.rows):
        cell = row.cells[0]
        cell_text = cell.text
        # 检查是否包含实际的周报内容（本周完成工作等）
        if '本周完成工作' in cell_text or '本周工作回顾' in cell_text:
            work_cell = cell
            work_row_idx = row_idx
            print(f"找到工作内容单元格: R{row_idx}")
            break
        # 如果是标题行，检查下一行
        if '工作复盘' in cell_text and row_idx + 1 < len(table.rows):
            next_cell = table.rows[row_idx + 1].cells[0]
            if '本周完成工作' in next_cell.text or len(next_cell.paragraphs) > 10:
                work_cell = next_cell
                work_row_idx = row_idx + 1
                print(f"找到工作内容单元格: R{row_idx + 1}")
                break
    
    if not work_cell:
        print("未找到工作内容单元格")
        return create_simple_report(organized_reports, output_path)
    
    # 分析单元格中的成员位置
    member_positions = analyze_work_content_cell(work_cell)
    print(f"找到 {len([k for k in member_positions if not k.startswith('__')])} 个成员位置")
    
    # 更新每个成员的内容
    updated_count = 0
    for member_name, content in organized_reports.items():
        if member_name in member_positions:
            start_idx = member_positions[member_name]
            if update_member_in_cell(work_cell, start_idx, content, member_positions):
                print(f"✓ 已更新: {member_name}")
                updated_count += 1
            else:
                print(f"△ 更新失败: {member_name}")
        else:
            print(f"✗ 模板中未找到: {member_name}")
    
    print(f"\n共更新 {updated_count} 人的周报")
    
    # === 3. 移除所有 XXX 占位符 ===
    xxx_removed = remove_xxx_placeholders_and_numbering(work_cell)
    if xxx_removed > 0:
        print(f"已移除 {xxx_removed} 个 XXX 占位符")
    
    # === 4. 清理空段落（保留成员之间的分隔空行）===
    member_names = set(member_positions.keys())
    empty_removed = clean_empty_paragraphs(work_cell, member_names)
    if empty_removed > 0:
        print(f"已清理 {empty_removed} 个空段落")
    
    # 保存
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    
    return output_path


def create_simple_report(
    organized_reports: Dict[str, Dict],
    output_path: str = None,
    week_info: str = "周报汇总"
) -> str:
    """
    创建简单格式的周报汇总（不依赖模板）
    """
    from content_parser import group_reports_by_team
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    output_path = output_path or config.OUTPUT_PATH
    
    doc = Document()
    
    # 标题
    title = doc.add_heading(week_info, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 按组分类
    grouped = group_reports_by_team(organized_reports)
    
    for group_name, members in config.GROUP_MEMBERS.items():
        group_reports = grouped.get(group_name, {})
        
        if not group_reports:
            continue
        
        doc.add_heading(group_name, level=1)
        
        for member_name in members:
            if member_name not in group_reports:
                continue
            
            content = group_reports[member_name]
            
            doc.add_heading(member_name, level=2)
            
            if content.get('this_week_work'):
                doc.add_heading('本周完成工作', level=3)
                doc.add_paragraph(content['this_week_work'])
            
            if content.get('next_week_plan'):
                doc.add_heading('下周工作计划', level=3)
                doc.add_paragraph(content['next_week_plan'])
            
            if content.get('gains_losses'):
                doc.add_heading('本周得与失', level=3)
                doc.add_paragraph(content['gains_losses'])
            
            doc.add_paragraph()
    
    # 未分组
    if grouped.get('未分组'):
        doc.add_heading('未分组', level=1)
        for member_name, content in grouped['未分组'].items():
            doc.add_heading(member_name, level=2)
            if content.get('raw_body'):
                doc.add_paragraph(content['raw_body'])
    
    doc.save(output_path)
    print(f"文档已保存到: {output_path}")
    
    return output_path


def create_report_document(
    organized_reports: Dict[str, Dict],
    template_path: str = None,
    output_path: str = None,
    week_info: str = None
) -> str:
    """
    创建周报汇总文档的主入口
    """
    return fill_template_by_replacement(
        organized_reports,
        template_path=template_path,
        output_path=output_path
    )


def read_template_structure(template_path: str) -> Dict:
    """
    读取模板结构（兼容旧接口）
    """
    doc = Document(template_path)
    
    structure = {
        'title': None,
        'praise_section_start': None,
        'suggestion_section_start': None,
        'work_section_start': None,
        'member_positions': {},
        'group_positions': {},
    }
    
    if doc.tables:
        table = doc.tables[0]
        for row_idx, row in enumerate(table.rows):
            cell_text = row.cells[0].text
            if '星光闪烁' in cell_text:
                structure['praise_section_start'] = row_idx
            if '诚信坦荡' in cell_text:
                structure['suggestion_section_start'] = row_idx
            if '工作复盘' in cell_text and not structure['work_section_start']:
                structure['work_section_start'] = row_idx
            
            # 如果当前行包含周报内容，分析成员位置
            if '本周完成工作' in cell_text or '本周工作回顾' in cell_text:
                member_pos = analyze_work_content_cell(row.cells[0])
                structure['member_positions'] = {
                    k: v for k, v in member_pos.items() 
                    if not k.startswith('__')
                }
    
    return structure


if __name__ == "__main__":
    print("测试模板分析...")
    
    try:
        template_path = find_template_file()
        print(f"模板文件: {template_path}")
        
        structure = read_template_structure(template_path)
        print(f"\n模板结构:")
        print(f"  星光闪烁: 行 {structure['praise_section_start']}")
        print(f"  诚信坦荡: 行 {structure['suggestion_section_start']}")
        print(f"  工作复盘: 行 {structure['work_section_start']}")
        print(f"  成员数: {len(structure['member_positions'])}")
        
        for name in list(structure['member_positions'].keys())[:5]:
            print(f"    - {name}")
        if len(structure['member_positions']) > 5:
            print(f"    ... 共 {len(structure['member_positions'])} 人")
            
    except Exception as e:
        print(f"错误: {e}")
